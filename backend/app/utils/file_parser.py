import pdfplumber
import docx
import os
from loguru import logger
from typing import Optional


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from a PDF file using pdfplumber.
    Returns extracted text or None if extraction fails.
    """
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    # Try extracting words if full text fails
                    words = page.extract_words()
                    if words:
                        text_parts.append(" ".join(w["text"] for w in words))
                logger.debug(f"Extracted text from PDF page {page_num + 1}")

        full_text = "\n".join(text_parts)
        return full_text.strip() if full_text.strip() else None

    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {e}")
        return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from a DOCX file using python-docx.
    Returns extracted text or None if extraction fails.
    """
    try:
        doc = docx.Document(file_path)
        text_parts = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        full_text = "\n".join(text_parts)
        logger.debug(f"Extracted {len(text_parts)} text blocks from DOCX")
        return full_text.strip() if full_text.strip() else None

    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
        return None


def extract_text(file_path: str, file_type: str) -> Optional[str]:
    """
    Extract text from either PDF or DOCX file.
    file_type should be 'pdf' or 'docx'.
    """
    file_type = file_type.lower().strip(".")
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file_path)
    else:
        logger.warning(f"Unsupported file type: {file_type}")
        return None


def validate_file_type(filename: str) -> tuple[bool, str]:
    """
    Validate that the file is a supported type (PDF or DOCX).
    Returns (is_valid, file_type).
    """
    ext = os.path.splitext(filename)[1].lower().strip(".")
    if ext == "pdf":
        return True, "pdf"
    elif ext in ("docx", "doc"):
        return True, "docx"
    return False, ""


def validate_file_size(file_size_bytes: int, max_size_mb: int = 10) -> bool:
    """Validate that file size is within the allowed limit."""
    max_size_bytes = max_size_mb * 1024 * 1024
    return file_size_bytes <= max_size_bytes
